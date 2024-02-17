from docx import Document
from docx2pdf import convert
import os
import PyPDF2
import sys
import datetime

# function untuk replace text
def replace_text(doc, old_text, new_text):
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old_text in cell.text:
                    cell.text = cell.text.replace(old_text, new_text)

# function untuk merge pdf dari application letter ke dokumen lain
def merge_pdfs(pdf_files, output_file):
    writer = PyPDF2.PdfWriter()

    for pdf_file in pdf_files:
        with open(pdf_file, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                writer.add_page(page)

    with open(output_file, 'wb') as output:
        writer.write(output)
        
# function generate tanggal saat ini
def gen_date():
    current_date = datetime.date.today()
    monthID = [
        "Januari", "Februari", "Maret", "April", "Mei", "Juni",
        "Juli", "Agustus", "September", "Oktober", "November", "Desember"
    ]
    # Mendapatkan nama bulan dari tanggal saat ini
    month_name = monthID[current_date.month - 1]

    # Format tanggal ke format yang diinginkan
    date_final = f"{current_date.day} {month_name} {current_date.year}"
    return date_final

def main():
    args = sys.argv[1:]
    company_wordlist = []
    job_wordlist = []
    location = ""

    # Input dari command line
    for i in range(len(args)):
        if args[i] == "-c" or args[i] == "--company":
            if os.path.exists(args[i+1]):
                with open(args[i+1], 'r', encoding='utf-8') as file:
                    company_wordlist = [line.strip() for line in file.readlines()]
            else:
                company_wordlist = [args[i+1]]
        elif args[i] == "-j" or args[i] == "--job-title":
            if os.path.exists(args[i+1]):
                with open(args[i+1], 'r', encoding='utf-8') as file:
                    job_wordlist = [line.strip() for line in file.readlines()]
            else:
                job_wordlist = [args[i+1]]
        elif args[i] == "-l" or args[i] == "--location":
            if os.path.exists(args[i+1]):
                with open(args[i+1], 'r', encoding='utf-8') as file:
                    location = file.read().strip()
            else:
                location = args[i+1]

    full_path = os.path.abspath(os.getcwd())

    # Load dokumen application leter
    doc_template_path = f"{full_path}\\docs\\SURAT_LAMARAN_KERJA.docx"
    # gabung templaet lokasi dan tanggal
    locdate=f'{location}, {gen_date()}'

    # Loop ke isi input dokumen
    for company_name, job_title in zip(company_wordlist, job_wordlist):
        # Load word untuk setiap iterasi
        doc = Document(doc_template_path)

        # replace template word
        replace_text(doc, "<<DATE>>", locdate)
        replace_text(doc, "<<PT. NAMA>>", company_name)
        replace_text(doc, "<<JOB>>", job_title)

        output_doc = f"{full_path}\\output\\Surat_Lamaran_Kerja_{company_name}_{job_title}.docx"
        doc.save(output_doc)

        # convert dokumen ke PDF
        output_pdf = f"{full_path}\\output\\pdf_output\\Surat_Lamaran_Kerja_{company_name}_{job_title}.pdf"
        convert(output_doc, output_pdf)
        
        # merge application letter pdf ke dokumen lain
        file_to_be_merged = f"{full_path}\\pdf-to-be-merge\\doc_to_merge.pdf"
        output_merged_pdf = f"{full_path}\\output\\pdf_merged_all\\Surat_Lamaran_Kerja_{company_name}_{job_title}.pdf"
        merge_pdfs([output_pdf, file_to_be_merged], output_merged_pdf)

        print(f"PDF for {company_name} - {job_title} saved successfully.")

if __name__ == '__main__':
    main()
